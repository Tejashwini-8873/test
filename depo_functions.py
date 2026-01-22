import os
import tempfile
import logging
from concurrent.futures import ThreadPoolExecutor
from PIL import Image
import base64
from openai import OpenAI
import requests
from docx import Document
import fitz  # PyMuPDF
import traceback
import re
import json
from docx.shared import Inches
import threading
import time
from datetime import datetime, timedelta
from dotenv import load_dotenv
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions

load_dotenv()

# --- Configuration & Constants ---

api_key = os.getenv("OPENAI_API_KEY")
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")

AZURE_STORAGE_CONNECTION_STRING = (
    "DefaultEndpointsProtocol=https;"
    "AccountName=depodatastorage;"
    "AccountKey=LyN82tPOGrvnh1nEReIzMj2jp5P6BMZZ2D4ypIFGNKqBcoWEAeic06AHrDBGUnjPBYs+gFoss4Ao+ASt6pUvtg==;"
    "EndpointSuffix=core.windows.net"
)

UPLOAD_CONTAINER = "depositions"
SUMMARY_CONTAINER = "summaries"

# Initialize Blob Service Client
try:
    blob_service = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
except Exception as e:
    logging.error(f"Failed to initialize BlobServiceClient: {e}")
    blob_service = None

# --- Text Extraction Functions ---

def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logging.error(f"Failed to extract text from DOCX: {str(e)}")
        raise

def extract_text_from_pdf(pdf_path: str) -> dict:
    """
    Extract text from a PDF file, returning a dictionary with page-wise text.
    
    Args:
        pdf_path (str): Path to the PDF file.
    
    Returns:
        dict: A dictionary where keys are page numbers and values are the extracted text.
    """
    import fitz
    import logging

    page_texts = {}
    try:
        pdf_document = fitz.open(pdf_path)
        for page_number in range(len(pdf_document)):
            page = pdf_document[page_number]
            text = page.get_text("text")
            page_texts[page_number + 1] = text.strip() if text else ""
        pdf_document.close()
        return page_texts
    except Exception as e:
        logging.error(f"Failed to extract text from PDF: {str(e)}")
        raise

# --- Azure Blob Storage Functions ---

def ensure_container(container_name):
    try:
        if blob_service:
            blob_service.create_container(container_name)
    except Exception:
        pass  # already exists

def upload_file_to_blob(uploaded_file):
    if not blob_service:
        raise Exception("Azure Blob Service not initialized.")

    blob_name = uploaded_file.name
    data = uploaded_file.getvalue()

    container = blob_service.get_container_client(UPLOAD_CONTAINER)
    blob = container.get_blob_client(blob_name)

    # Upload blob
    blob.upload_blob(data, overwrite=True)
    logging.info(f"📤 Deposition uploaded: {blob_name}")

    # Generate SAS URL so user can view the deposition
    sas_token = generate_blob_sas(
        account_name=blob_service.account_name,
        container_name=UPLOAD_CONTAINER,
        blob_name=blob_name,
        account_key=blob_service.credential.account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=24)
    )

    sas_url = (
        f"https://{blob_service.account_name}.blob.core.windows.net/"
        f"{UPLOAD_CONTAINER}/{blob_name}?{sas_token}"
    )

    logging.info(f"🔐 Deposition SAS URL generated: {sas_url}")

    return blob_name, sas_url

def download_blob_to_temp(blob_name):
    """
    Downloads blob → returns temporary local file path for processing.
    """
    if not blob_service:
        raise Exception("Azure Blob Service not initialized.")
        
    container = blob_service.get_container_client(UPLOAD_CONTAINER)
    blob = container.get_blob_client(blob_name)

    data = blob.download_blob().readall()

    suffix = ".pdf" if blob_name.lower().endswith(".pdf") else ".docx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        return tmp.name

def upload_summary_to_blob(local_path, new_blob_name):
    if not blob_service:
        raise Exception("Azure Blob Service not initialized.")

    container = blob_service.get_container_client(SUMMARY_CONTAINER)
    blob = container.get_blob_client(new_blob_name)

    # Upload file
    with open(local_path, "rb") as f:
        blob.upload_blob(f, overwrite=True)

    logging.info(f"📤 Summary uploaded successfully as '{new_blob_name}'")

    # Generate SAS token (read-only, 24 hours)
    sas_token = generate_blob_sas(
        account_name=blob_service.account_name,
        container_name=SUMMARY_CONTAINER,
        blob_name=new_blob_name,
        account_key=blob_service.credential.account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=24)
    )

    sas_url = (
        f"https://{blob_service.account_name}.blob.core.windows.net/"
        f"{SUMMARY_CONTAINER}/{new_blob_name}?{sas_token}"
    )

    logging.info(f"🔐 SAS URL generated: {sas_url}")
    return sas_url

# --- OpenAI Interaction ---

def get_chatgpt_response(prompt: str,text: str, api_key: str, model: str) -> str:
    """
    Generate a summary using OpenAI ChatGPT API.
    """
    client = OpenAI(api_key=api_key)
    response = client.responses.create(
        model=model,
        input=[
           {"role": "system", "content": "You are a senior legal deposition summarizer. Use the provided instructions and the deposition text to generate a clear, concise summary "},
        {"role": "user", "content": f"{prompt}\n\nContract Text:\n{text}"}
        ]
    )

    summary_text = response.output_text
    return summary_text

# def get_chatgpt_response(prompt: str, text: str, api_key: str, model: str = "gpt-5") -> str:
#     """
#     Generate a summary using OpenAI ChatGPT API.
#     """
#     client = OpenAI(api_key=api_key)
#     response = client.responses.create(
#         model=model,
#         input=[
#            {"role": "system", "content": "You are a senior legal deposition summarizer. Use the provided instructions and the deposition text to generate a clear, concise summary "},
#         {"role": "user", "content": f"{prompt}\n\nContract Text:\n{text}"}
#         ]
#     )

#     summary_text = response.output_text
#     return summary_text

# --- JSON & Parsing Helpers ---

def extract_page_group_json(text: str):
    """
    Extracts Page-Group Subject Summaries JSON from GPT output.
    """
    # 1️⃣ Remove known headers / labels
    cleaned = re.sub(
        r"Page-Group Subject Summaries\s*\(JSON\)\s*",
        "",
        text,
        flags=re.IGNORECASE
    ).strip()

    # 2️⃣ Try direct JSON parse first
    try:
        data = json.loads(cleaned)
        if isinstance(data, dict):
            return [data]
        if isinstance(data, list):
            return data
    except Exception:
        pass

    # 3️⃣ Extract JSON array if present
    array_match = re.search(r"\[\s*\{.*?\}\s*\]", cleaned, re.DOTALL)
    if array_match:
        return json.loads(array_match.group(0))

    # 4️⃣ Extract multiple standalone JSON objects
    objects = []
    for match in re.finditer(r"\{.*?\}", cleaned, re.DOTALL):
        try:
            obj = json.loads(match.group(0))
            if isinstance(obj, dict) and "subject" in obj:
                objects.append(obj)
        except Exception:
            continue

    if objects:
        return objects

    # 5️⃣ Nothing valid found
    raise ValueError("No valid Page-Group JSON found in GPT output")

def extract_exhibits_table(text):
    """Extract markdown-style 'Exhibits Table' into structured rows"""
    pattern = r"(?s)Exhibits Table\s*\|.*?\|\n(.*?)\n(?:\s*\n|$)"
    match = re.search(pattern, text)
    if not match:
        return None

    table_text = match.group(1).strip()
    rows = []
    for line in table_text.split("\n"):
        line = line.strip()
        if not line or line.startswith("|---"):
            continue

        cols = [col.strip() for col in line.strip("|").split("|")]
        if len(cols) >= 3:
            rows.append({
                "Exhibit No./Name": cols[0],
                "Page Numbers": cols[1],
                "Brief Description & Relevance": cols[2]
            })
    return rows

# --- Document Generation ---

def create_deposition_summary(input_docx, output_docx):
    # Load input DOCX
    doc = Document(input_docx)
    text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    # ------------------------------------
    # STEP 1: Extract JSON from the text
    # ------------------------------------
    try:
        # Locate the section between the markers
        start_marker = "Page-Group Subject Summaries"
        end_marker = "Structured Deposition Summary"

        start_index = text.find(start_marker)
        end_index = text.find(end_marker)

        if start_index == -1 or end_index == -1:
            raise ValueError("Could not locate start or end markers in document text.")

        # Extract the portion between the markers
        json_block = text[start_index + len(start_marker):end_index].strip()

        # Remove everything before first [ and after last ]
        json_start = json_block.find("[")
        json_end = json_block.rfind("]") + 1
        json_str = json_block[json_start:json_end]

        # Clean up any fancy dashes or non-breaking spaces
        json_str = json_str.replace("–", "-").replace("\u2013", "-").replace("\u00a0", " ")
        
        # Parse JSON
        deposition_data = extract_page_group_json(json_str)
        print("✅ Extracted deposition data from JSON successfully.")

    except Exception as e:
        print(f"⚠️ Failed to extract JSON: {e}")
        snippet = text[start_index:start_index + 300] if start_index != -1 else text[:300]
        print("📄 Extracted text snippet for debugging:\n", snippet)
        return

    # ------------------------------------
    # STEP 2: Create new formatted DOCX
    # ------------------------------------
    out = Document()
    out.add_heading("Deposition Summary", level=1)

    # -----------------------------
    # SECTION 1: Page-Group Table
    # -----------------------------
    out.add_paragraph("")
    out.add_heading("1. Page-Group Subject Summaries", level=2)
    out.add_paragraph("")

    table = out.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Subject & Content"
    hdr_cells[1].text = "Page:Line Range"

    # Fill table
    for entry in deposition_data:
        line_refs = []
        for page, lines in entry["line_numbers"].items():
            if lines:
                sorted_lines = sorted(lines)
                if len(sorted_lines) == 1:
                    line_refs.append(f"{page}:{sorted_lines[0]}")
                else:
                    line_refs.append(f"{page}:{sorted_lines[0]}-{sorted_lines[-1]}")
        line_str = "\n".join(line_refs)

        row_cells = table.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run(entry["subject"] + "; ")
        run.bold = True
        paragraph.add_run(entry["content"])
        row_cells[1].text = line_str

    # -----------------------------
    # STEP 3: Structured Summary
    # -----------------------------
    out.add_paragraph("")
    out.add_paragraph("")

    # Extract structured summary text (everything after JSON block)
    structured_text = text[end_index:].strip()

    # Split sections by numbered headers (like "1. Legal Issue", "2. Purpose ...")
    sections = re.split(r"(?=\n\d+\.\s)", "\n" + structured_text)
    sections = [s.strip() for s in sections if s.strip()]

    for section in sections:
        lines = section.splitlines()
        header = lines[0].strip()
        out.add_heading(header, level=3)
        content = "\n".join(lines[1:]).strip()

         # ---- Exhibit Table Auto Extraction ----
        if "Exhibits Table" in header:
            # Capture lines with '|' that look like a table
            table_lines = [l for l in content.splitlines() if "|" in l and not l.startswith("|---")]
            if table_lines:
                # Parse header and rows
                header_row = [h.strip() for h in table_lines[0].strip("|").split("|")]
                data_rows = []
                for row in table_lines[1:]:
                    cols = [c.strip() for c in row.strip("|").split("|")]
                    if len(cols) >= len(header_row):
                        data_rows.append(cols[:len(header_row)])

                # Create Word table
                exhibit_table = out.add_table(rows=1, cols=len(header_row))
                exhibit_table.style = 'Table Grid'
                hdr_cells = exhibit_table.rows[0].cells
                for i, col_name in enumerate(header_row):
                    hdr_cells[i].text = col_name

                for row in data_rows:
                    row_cells = exhibit_table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = val
                continue
        # --------------------------------------

        # Add normal section text or bullet points
        for para in content.split("\n"):
            if para.strip().startswith("- "):
                out.add_paragraph(para.strip("- ").strip(), style='List Bullet')
            elif para.strip():
                out.add_paragraph(para.strip())

    out.save(output_docx)
    print(f"✅ Formatted output saved as: {output_docx}")

def save_as_docx(summary, filename):
    doc = Document()
    doc.add_heading("Deposition Summary", level=1)
    doc.add_paragraph(summary)
    output_path = os.path.join(tempfile.gettempdir(), f"{filename}.docx")
    doc.save(output_path)
    return output_path

# --- Background Worker ---

def background_summary(blob_name, api_key_arg, prompt_text):
    # If api_key not passed, try global
    if not api_key_arg:
        api_key_arg = api_key
        
    logs = []

    def log(msg):
        # central logging for both console and return payload
        timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
        entry = f"{timestamp} - {msg}"
        print(entry)
        logs.append(entry)

    try:
        log("background_summary() STARTED")
        log(f"Blob requested: {blob_name}")

        # 1️⃣ Download deposition file from Azure Blob
        log("📥 Downloading from Azure Blob...")
        temp_path = download_blob_to_temp(blob_name)
        log(f"Downloaded to temp path: {temp_path}")

        # 2️⃣ Extract text
        if temp_path.lower().endswith(".pdf"):
            log("🧾 Detected PDF — extracting text from PDF...")
            text = extract_text_from_pdf(temp_path)
            log(f"PDF extraction complete — extracted {len(text)} characters")
        else:
            log("📝 Detected DOCX — extracting text from DOCX...")
            text = extract_text_from_docx(temp_path)
            log(f"DOCX extraction complete — extracted {len(text)} characters")

        # 3️⃣ Generate summary using GPT
        log("🤖 Calling get_chatgpt_response() to generate summary...")
        try:
            summary_text = get_chatgpt_response(prompt_text, text, api_key_arg, model="gpt-5")
            log(f"AI summary generated — length {len(summary_text)} characters")
        except Exception as e:
            err = traceback.format_exc()
            log(f"❌ get_chatgpt_response() failed: {e}")
            log(err)
            raise

        # 4️⃣ Save raw summary locally for formatting
        base_name = os.path.splitext(blob_name)[0]
        raw_local = os.path.join(tempfile.gettempdir(), f"{base_name}_summary_raw.docx")
        log(f"Saving raw summary to: {raw_local}")
        doc = Document()
        doc.add_heading("Deposition Summary", level=1)
        doc.add_paragraph(summary_text)
        doc.save(raw_local)
        log("Raw summary saved.")

        # 5️⃣ Try formatted summary
        final_local = os.path.join(tempfile.gettempdir(), f"{base_name}_summary_final.docx")
        try:
            log("Attempting to format raw summary into structured DOCX...")
            create_deposition_summary(raw_local, final_local)
            if not os.path.exists(final_local):
                raise ValueError("Formatted DOCX was not created.")
            final_used = final_local
            log("✅ Formatting applied successfully.")
        except Exception as e:
            fmt_err = traceback.format_exc()
            log(f"⚠️ Formatting failed: {e}")
            log(fmt_err)
            final_used = raw_local
            log("Using raw summary as fallback.")

        # 6️⃣ Upload summary DOCX to Azure Blob Storage
        final_blob_name = f"{base_name}_summary.docx"
        log(f"Uploading final summary to Azure as blob: {final_blob_name}")
        final_url = upload_summary_to_blob(final_used, final_blob_name)
        log(f"🚀 Uploaded summary to Azure Blob Storage: {final_url}")

        log("background_summary() COMPLETED")
        return {"path": final_url, "log": logs}

    except Exception as e:
        err_trace = traceback.format_exc()
        log(f"FATAL ERROR in background_summary(): {e}")
        log(err_trace)
        return {"path": None, "error": str(e), "log": logs}

# --- Utils ---

def get_base64_image(image_url):
    response = requests.get(image_url)
    return base64.b64encode(response.content).decode()
