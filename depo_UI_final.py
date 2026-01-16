import os
import streamlit as st
import tempfile
import logging
from concurrent.futures import ThreadPoolExecutor
from PIL import Image
import base64  
import os
import streamlit as st
import openai
import requests
from docx import Document
import fitz
import tempfile
import traceback
import tempfile
from depo_gpt5 import extract_text_from_docx, extract_text_from_pdf, generate_summary_with_chatgpt
import re
import json
from docx.shared import Inches
from concurrent.futures import ThreadPoolExecutor
import threading
import time
import base64
import requests
from streamlit_autorefresh import st_autorefresh


# ------------------ AZURE BLOB STORAGE SETUP ------------------
from azure.storage.blob import BlobServiceClient

from azure.storage.blob import generate_blob_sas, BlobSasPermissions
from datetime import datetime, timedelta

from dotenv import load_dotenv
load_dotenv()


# --- Logging ---
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Ensure session state keys exist
st.session_state.setdefault("summary_status", "idle")
st.session_state.setdefault("summary_future", None)
st.session_state.setdefault("summary_result", None)
st.session_state.setdefault("summary_log", [])
st.session_state.setdefault("summary_error", None)
st.session_state.setdefault("summary_completed_once", False)
st.session_state.setdefault("summary_needs_rerun", False)
st.session_state.setdefault("show_status_msg", False)
st.session_state.setdefault("pause_autorefresh", False)



import time

if (
    st.session_state.summary_status == "running"
    and not st.session_state.pause_autorefresh
):
    st_autorefresh(interval=10000, key="summary_poll")



if st.session_state.summary_status == "running":
    future = st.session_state.get("summary_future")

    if (
    st.session_state.summary_status == "running"
    and future
    and future.done()):
        result = future.result()
        

        st.session_state.summary_log = result.get("log", [])

        if result.get("path"):
            st.session_state.summary_result = result["path"]
            st.session_state.summary_status = "done"
            
            st.session_state.show_status_msg = False   
        else:
            st.session_state.summary_error = result.get("error", "Unknown error")
            st.session_state.summary_status = "error"
            st.session_state.show_status_msg = False  
        st.session_state.summary_completed_once = True
        

if st.session_state.summary_needs_rerun:
    st.session_state.summary_needs_rerun = False
    st.rerun()

if st.session_state.summary_completed_once:
    st.session_state.summary_completed_once = False

    st.rerun()


        # st.rerun()

# # --- API Keys (use env vars for production) ---

api_key = os.getenv("OPENAI_API_KEY")
PERPLEXITY_API_KEY = os.getenv("PERPLEXITY_API_KEY")

# YOUR STORAGE CONNECTION STRING
AZURE_STORAGE_CONNECTION_STRING = (
    "DefaultEndpointsProtocol=https;"
    "AccountName=depodatastorage;"
    "AccountKey=LyN82tPOGrvnh1nEReIzMj2jp5P6BMZZ2D4ypIFGNKqBcoWEAeic06AHrDBGUnjPBYs+gFoss4Ao+ASt6pUvtg==;"
    "EndpointSuffix=core.windows.net"
)

# Create Blob Client
blob_service = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)

# Containers used
UPLOAD_CONTAINER = "depositions"
SUMMARY_CONTAINER = "summaries"

def upload_file_to_blob(uploaded_file):
    blob_name = uploaded_file.name
    data = uploaded_file.getvalue()

    container = blob_service.get_container_client(UPLOAD_CONTAINER)
    blob = container.get_blob_client(blob_name)

    # Upload blob
    blob.upload_blob(data, overwrite=True)
    logging.info(f"📤 Deposition uploaded: {blob_name}")

    # Generate SAS URL so user can view the deposition
    sas_token = generate_blob_sas(
        account_name=blob_service.account_name,
        container_name=UPLOAD_CONTAINER,
        blob_name=blob_name,
        account_key=blob_service.credential.account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=24)
    )

    sas_url = (
        f"https://{blob_service.account_name}.blob.core.windows.net/"
        f"{UPLOAD_CONTAINER}/{blob_name}?{sas_token}"
    )

    logging.info(f"🔐 Deposition SAS URL generated: {sas_url}")

    return blob_name, sas_url



def download_blob_to_temp(blob_name):
    """
    Downloads blob → returns temporary local file path for processing.
    """
    container = blob_service.get_container_client(UPLOAD_CONTAINER)
    blob = container.get_blob_client(blob_name)

    data = blob.download_blob().readall()

    suffix = ".pdf" if blob_name.lower().endswith(".pdf") else ".docx"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        return tmp.name
def upload_summary_to_blob(local_path, new_blob_name):
    container = blob_service.get_container_client(SUMMARY_CONTAINER)
    blob = container.get_blob_client(new_blob_name)

    # Upload file
    with open(local_path, "rb") as f:
        blob.upload_blob(f, overwrite=True)

    logging.info(f"📤 Summary uploaded successfully as '{new_blob_name}'")

    # Generate SAS token (read-only, 24 hours)
    sas_token = generate_blob_sas(
        account_name=blob_service.account_name,
        container_name=SUMMARY_CONTAINER,
        blob_name=new_blob_name,
        account_key=blob_service.credential.account_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=24)
    )

    sas_url = (
        f"https://{blob_service.account_name}.blob.core.windows.net/"
        f"{SUMMARY_CONTAINER}/{new_blob_name}?{sas_token}"
    )

    logging.info(f"🔐 SAS URL generated: {sas_url}")
    return sas_url

# Ensure containers exist
def ensure_container(container_name):
    try:
        blob_service.create_container(container_name)
    except Exception:
        pass  # already exists


# ensure_container(UPLOAD_CONTAINER)
# ensure_container(SUMMARY_CONTAINER)


executor = ThreadPoolExecutor(max_workers=1)
def create_deposition_summary(input_docx, output_docx):
    # Load input DOCX
    doc = Document(input_docx)
    text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    # ------------------------------------
    # STEP 1: Extract JSON from the text
    # ------------------------------------
    try:
        # Locate the section between the markers
        start_marker = "Page-Group Subject Summaries"
        end_marker = "Structured Deposition Summary"

        start_index = text.find(start_marker)
        end_index = text.find(end_marker)

        if start_index == -1 or end_index == -1:
            raise ValueError("Could not locate start or end markers in document text.")

        # Extract the portion between the markers
        json_block = text[start_index + len(start_marker):end_index].strip()

        # Remove everything before first [ and after last ]
        json_start = json_block.find("[")
        json_end = json_block.rfind("]") + 1
        json_str = json_block[json_start:json_end]

        # Clean up any fancy dashes or non-breaking spaces
        json_str = json_str.replace("–", "-").replace("\u2013", "-").replace("\u00a0", " ")
        # print(("Cleaned JSON String:", json_str))
        # Parse JSON
        deposition_data = json.loads(json_str)
        print("✅ Extracted deposition data from JSON successfully.")

    except Exception as e:
        print(f"⚠️ Failed to extract JSON: {e}")
        snippet = text[start_index:start_index + 300] if start_index != -1 else text[:300]
        print("📄 Extracted text snippet for debugging:\n", snippet)
        return

    # ------------------------------------
    # STEP 2: Create new formatted DOCX
    # ------------------------------------
    out = Document()
    out.add_heading("Deposition Summary", level=1)

    # -----------------------------
    # SECTION 1: Page-Group Table
    # -----------------------------
    out.add_paragraph("")
    out.add_heading("1. Page-Group Subject Summaries", level=2)
    out.add_paragraph("")

    table = out.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Subject & Content"
    hdr_cells[1].text = "Page:Line Range"

    # Fill table
    for entry in deposition_data:
        line_refs = []
        for page, lines in entry["line_numbers"].items():
            if lines:
                sorted_lines = sorted(lines)
                if len(sorted_lines) == 1:
                    line_refs.append(f"{page}:{sorted_lines[0]}")
                else:
                    line_refs.append(f"{page}:{sorted_lines[0]}-{sorted_lines[-1]}")
        line_str = "\n".join(line_refs)

        row_cells = table.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run(entry["subject"] + "; ")
        run.bold = True
        paragraph.add_run(entry["content"])
        row_cells[1].text = line_str

    # -----------------------------
    # STEP 3: Structured Summary
    # -----------------------------
    out.add_paragraph("")
    # out.add_heading("2. Structured Deposition Summary", level=2)
    out.add_paragraph("")

    # Extract structured summary text (everything after JSON block)
    structured_text = text[end_index:].strip()

    # Split sections by numbered headers (like "1. Legal Issue", "2. Purpose ...")
    sections = re.split(r"(?=\n\d+\.\s)", "\n" + structured_text)
    sections = [s.strip() for s in sections if s.strip()]

    for section in sections:
        lines = section.splitlines()
        header = lines[0].strip()
        out.add_heading(header, level=3)
        content = "\n".join(lines[1:]).strip()

         # ---- Exhibit Table Auto Extraction ----
        if "Exhibits Table" in header:
            # Capture lines with '|' that look like a table
            table_lines = [l for l in content.splitlines() if "|" in l and not l.startswith("|---")]
            if table_lines:
                # Parse header and rows
                header_row = [h.strip() for h in table_lines[0].strip("|").split("|")]
                data_rows = []
                for row in table_lines[1:]:
                    cols = [c.strip() for c in row.strip("|").split("|")]
                    if len(cols) >= len(header_row):
                        data_rows.append(cols[:len(header_row)])

                # Create Word table
                exhibit_table = out.add_table(rows=1, cols=len(header_row))
                exhibit_table.style = 'Table Grid'
                hdr_cells = exhibit_table.rows[0].cells
                for i, col_name in enumerate(header_row):
                    hdr_cells[i].text = col_name

                for row in data_rows:
                    row_cells = exhibit_table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = val
                continue
        # --------------------------------------


        # Add normal section text or bullet points
        for para in content.split("\n"):
            if para.strip().startswith("- "):
                out.add_paragraph(para.strip("- ").strip(), style='List Bullet')
            elif para.strip():
                out.add_paragraph(para.strip())

    out.save(output_docx)
    print(f"✅ Formatted output saved as: {output_docx}")
def extract_exhibits_table(text):
    """Extract markdown-style 'Exhibits Table' into structured rows"""
    pattern = r"(?s)Exhibits Table\s*\|.*?\|\n(.*?)\n(?:\s*\n|$)"
    match = re.search(pattern, text)
    if not match:
        return None

    table_text = match.group(1).strip()
    rows = []
    for line in table_text.split("\n"):
        line = line.strip()
        if not line or line.startswith("|---"):
            continue

        cols = [col.strip() for col in line.strip("|").split("|")]
        if len(cols) >= 3:
            rows.append({
                "Exhibit No./Name": cols[0],
                "Page Numbers": cols[1],
                "Brief Description & Relevance": cols[2]
            })
    return rows    



def save_as_docx(summary, filename):
    doc = Document()
    doc.add_heading("Deposition Summary", level=1)
    doc.add_paragraph(summary)
    output_path = os.path.join(tempfile.gettempdir(), f"{filename}.docx")
    doc.save(output_path)
    return output_path

def save_uploaded_file(uploaded_file):
    # 🔧 Hardcode output paths for debugging
    base_name = os.path.splitext(uploaded_file.name)[0]
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path
# ...existing code...
def background_summary(blob_name, api_key, prompt_text):
    logs = []

    def log(msg):
        # central logging for both console and return payload
        timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
        entry = f"{timestamp} - {msg}"
        print(entry)
        logs.append(entry)

    try:
        log("background_summary() STARTED")
        log(f"Blob requested: {blob_name}")

        # 1️⃣ Download deposition file from Azure Blob
        log("📥 Downloading from Azure Blob...")
        temp_path = download_blob_to_temp(blob_name)
        log(f"Downloaded to temp path: {temp_path}")

        # 2️⃣ Extract text
        if temp_path.lower().endswith(".pdf"):
            log("🧾 Detected PDF — extracting text from PDF...")
            text = extract_text_from_pdf(temp_path)
            log(f"PDF extraction complete — extracted {len(text)} characters")
        else:
            log("📝 Detected DOCX — extracting text from DOCX...")
            text = extract_text_from_docx(temp_path)
            log(f"DOCX extraction complete — extracted {len(text)} characters")

        # 3️⃣ Generate summary using GPT
        log("🤖 Calling get_chatgpt_response() to generate summary...")
        try:
            summary_text = get_chatgpt_response(prompt_text, text, api_key, model="gpt-5")
            log(f"AI summary generated — length {len(summary_text)} characters")
        except Exception as e:
            err = traceback.format_exc()
            log(f"❌ get_chatgpt_response() failed: {e}")
            log(err)
            raise

        # 4️⃣ Save raw summary locally for formatting
        base_name = os.path.splitext(blob_name)[0]
        raw_local = os.path.join(tempfile.gettempdir(), f"{base_name}_summary_raw.docx")
        log(f"Saving raw summary to: {raw_local}")
        doc = Document()
        doc.add_heading("Deposition Summary", level=1)
        doc.add_paragraph(summary_text)
        doc.save(raw_local)
        log("Raw summary saved.")

        # 5️⃣ Try formatted summary
        final_local = os.path.join(tempfile.gettempdir(), f"{base_name}_summary_final.docx")
        try:
            log("Attempting to format raw summary into structured DOCX...")
            create_deposition_summary(raw_local, final_local)
            final_used = final_local
            log("✅ Formatting applied successfully.")
        except Exception as e:
            fmt_err = traceback.format_exc()
            log(f"⚠️ Formatting failed: {e}")
            log(fmt_err)
            final_used = raw_local
            log("Using raw summary as fallback.")

        # 6️⃣ Upload summary DOCX to Azure Blob Storage
        final_blob_name = f"{base_name}_summary.docx"
        log(f"Uploading final summary to Azure as blob: {final_blob_name}")
        final_url = upload_summary_to_blob(final_used, final_blob_name)
        log(f"🚀 Uploaded summary to Azure Blob Storage: {final_url}")

        log("background_summary() COMPLETED")
        return {"path": final_url, "log": logs}

    except Exception as e:
        err_trace = traceback.format_exc()
        log(f"FATAL ERROR in background_summary(): {e}")
        log(err_trace)
        return {"path": None, "error": str(e), "log": logs}

def get_base64_image(image_url):
    response = requests.get(image_url)
    return base64.b64encode(response.content).decode()

def get_chatgpt_response(prompt: str,text: str, api_key: str, model: str) -> str:
    """
    Generate a summary using OpenAI ChatGPT API.
    """
    import openai
    openai.api_key = api_key
    messages = [
        {"role": "system", "content": "You are a senior legal deposition summarizer. Use the provided instructions and the deposition text to generate a clear, concise summary "},
        {"role": "user", "content": f"{prompt}\n\nContract Text:\n{text}"}
    ]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        request_timeout=1200,   #  Increase timeout
        timeout=1200
        # reasoning={"effort": "high"}
    )
    return response["choices"][0]["message"]["content"].strip()

# --- Page Config ---
st.set_page_config(page_title="Deposition Summarizer", page_icon="🧾", layout="wide")


# --- Config ---
st.set_page_config(page_title="RLG | Depo Summarizer ", page_icon="📜", layout="wide")
executor = ThreadPoolExecutor(max_workers=1)
st.markdown("""
    <style>
        .main, body, [class*="block-container"] {
            background-color: #f2fbf5 !important;  /* 🌿 light mint green */
        }
    </style>
""", unsafe_allow_html=True)

# --- Custom CSS ---
st.markdown("""
    <style>
        .main {
            background-color: #f5f8f5;
        }
        

        /* Top Bar */
        .top-bar {
            display: flex;
            align-items: center;
            justify-content: flex-start;
            padding: 0.6rem 1.2rem;
            background: linear-gradient(90deg, #009e60, #00b26d);
            color: white;
            border-bottom: 2px solid #007c48;
        }
        .logo {
            height: 50px;
            margin-right: 14px;
            border-radius: 6px;
            box-shadow: 0 0 6px rgba(0,0,0,0.1);
        }
        .title {
            font-size: 1.9rem;
            font-weight: 800;
            letter-spacing: -0.5px;
        }

        /* Main Title */
        .main-title {
            text-align: center;
            font-weight: 900;
            font-size: 2.1em;
            color: #007c48;
            margin-bottom: 10px;
        }

        /* Buttons */
        .stButton>button {
            background-color: #009e60;
            color: white;
            font-weight: 600;
            border-radius: 10px;
            padding: 0.6em 1.3em;
            border: none;
            transition: all 0.3s ease;
            box-shadow: 0 2px 5px rgba(0,158,96,0.2);
        }
        .stButton>button:hover {
            background-color: #00b26d;
            transform: scale(1.03);
            box-shadow: 0 5px 12px rgba(0,158,96,0.25);
        }

        /* Sidebar */
        # section[data-testid="stSidebar"] {
        #     background: #f1f8f2;
        #     border-right: 2px solid #b6dec2;
        # }
        .sidebar-title {
            font-weight: 700;
            color: #009e60;
            margin-bottom: 0.5rem;
            font-size: 1.1em;
        }
        .chat-item {
            background: #ffffff;
            border-radius: 10px;
            padding: 0.6rem;
            margin-bottom: 0.5rem;
            box-shadow: 0 2px 6px rgba(0,0,0,0.05);
            cursor: pointer;
            transition: 0.2s;
        }
        .chat-item:hover {
            background: #e5f7ec;
        }

        /* Response Cards */
        .response-box {
            background-color: #ffffff;
            border-left: 6px solid #009e60;
            border-radius: 15px;
            padding: 1em 1.3em;
            box-shadow: 0 3px 8px rgba(0,0,0,0.06);
            margin-bottom: 1.2em;
        }
        .question {
            font-weight: 600;
            color: #006b3f;
        }
        .answer {
            background: #f2fbf5;
            border-left: 4px solid #00b26d;
            padding: 0.8em 1em;
            border-radius: 8px;
            margin-top: 0.4em;
        }

        /* Footer */
        .footer {
            text-align: center;
            color: #444;
            font-size: 0.9em;
            padding-top: 20px;
            border-top: 1px solid #b6dec2;
            margin-top: 35px;
        }
    </style>
""", unsafe_allow_html=True)
st.markdown("""
<style>

/* 1️⃣ Push sidebar DOWN below the 140px header */
section[data-testid="stSidebar"] {
    margin-top: 150px !important;
}



/* 3️⃣ The container wrapping the toggle also needs to be pushed down */
div[data-testid="collapsedControl"] {
    position: fixed !important;
    top: 150px !important;
    left: 10px !important;
    z-index: 5005 !important;
}

/* 4️⃣ Remove Streamlit’s default shadow that overlaps header */
header {
    position: relative !important;
    z-index: 1 !important;
}

</style>
""", unsafe_allow_html=True)


# logo_path = r"C:\Users\Teju\Downloads\twc.webp"
logo_path = r"https://raw.githubusercontent.com/Tejashwini-8873/test/main/assets/RLG.jpg"
logo_base64 = get_base64_image(logo_path)




# --- Render Compact Green Header ---
st.markdown(f"""
    <style>
        .top-header {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 140px; /* 🌿 reduced height */
            background: linear-gradient(90deg, rgba(0,158,96,0.9), rgba(0,178,109,0.92)),
                        url("data:image/webp;base64,{logo_base64}") no-repeat left center;
            background-size: auto 140px; /* scale logo to new height */
            background-blend-mode: overlay;
            border-bottom: 3px solid #007c48;
            box-shadow: 0 3px 8px rgba(0,0,0,0.2);
            display: flex;
            flex-direction: column;
            justify-content: center;
            align-items: center;
            text-align: center;
            z-index: 1000;
        }}

        .top-header h1 {{
            color: white;
            font-size: 2rem; /* slightly smaller title */
            font-weight: 900;
            margin: 0 0 6px 0;
            letter-spacing: -0.5px;
        }}

        .top-header p {{
            color: #e8ffe9;
            font-size: 1rem;
            font-style: italic;
            margin: 0;
        }}

        /* Adjust page padding below header */
        .block-container {{
            padding-top: 120px !important;
        }}
    </style>

    <div class="top-header">
        <h1>📜 RLG Deposition Summarizer</h1>
        <p>AI-powered legal deposition analysis — with The Wonderful touch 🍃</p>
    </div>
""", unsafe_allow_html=True)

# --- Session State ---
if "user_responses" not in st.session_state:
    st.session_state["user_responses"] = []
if "selected_chat_index" not in st.session_state:
    st.session_state["selected_chat_index"] = None


uploaded_file = st.file_uploader("📂 Upload a deposition document (PDF or Word) and let AI  summarize and extract key legal insights effortlessly.", type=["pdf", "docx"])

# Save file temporarily
def save_uploaded_file(uploaded_file):
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path

col1, col2, col3 = st.columns([1, 1, 1])
# prompt = "Summarize the deposition in 2 points"




json_format= """{
            "subject": "<short header summarizing the topic of these pages>",
            "content": "<1-2 line factual mini-summary of the testimony or events in these pages>",
            "line_numbers": {
            "<page_number>": [<only the most relevant line numbers from this page>],
            "<page_number>": [<only the most relevant line numbers from this page>]
            }
        },
        {
            "subject": "<short header summarizing the topic of these pages>",
            "content": "<1-2 line factual mini-summary of the testimony or events in these pages>",
            "line_numbers": {
            "<page_number>": [<only the most relevant line numbers from this page>],
            "<page_number>": [<only the most relevant line numbers from this page>]
            }
        }
        """


prompt = f"""
        You are a senior legal analyst specializing in deposition analysis. Your task is to review a full deposition transcript and perform two critical functions:
        1. Page-Group Subject Summaries
        2. A structured, professional legal summary

        ---
        
        # ### 1. Page-Group Subject Summaries
            - You must review the entire deposition transcript thoroughly from start to end. 
            - Divide the transcript into sequential, non-overlapping chunks. Each chunk must:
                • Cover a continuous range of pages in order.  
                • Group together related discussions, testimony, or objections that form a coherent subject.  
                • Ensure that every page of the transcript is included in exactly one chunk (no page may be skipped or left out).  
            - For each chunk:
                • Identify a concise subject line summarizing the primary topic or testimony for that page range.  
                • Write a 2–3 line factual mini-summary of the testimony in that chunk.  
                • Keep the summary neutral, objective, and legally relevant (no opinions or speculation).  
                - Provide a **2–3 line factual mini-summary** of the content in those pages make sure all pages are included.
                - Keep it **neutral, objective, and legally relevant**.
                • Ensure that every page of the transcript is considered in sequence, but you may omit chunks if:
                    – The pages contain no substantive facts to summarize (e.g., filler, procedural headers , word glossary ).  
                    – The pages have no valid line numbers available for extraction
            - Output must cover the **entire deposition**, from the first page to the last, in properly ordered chunks.
                            
            - **Line_Numbers**:  
                - Never invent or guess line numbers. Use only numbers that truly appear on that page (available_lines or parsed from page_text).
                - Parsing: Each page is already provided as a dictionary (`1: "...", 2: "...", ...`).   — i.e., a dict mapping **page → list of 4–5 line numbers**.
                Inside each page, every `\n` corresponds to a new line number, which is marked at the start (e.g., `1`, `2`, `3`).  
                This allows you to map: **page → line numbers → text** cleanly.  
                - Select only line numbers that directly support the chunk’s subject/summary.
                - Do not include every line or filler text (e.g., "Page X", "Veritext Legal Solutions").
                - Output as a dict: "page": [line1, line2, ...].
                - Each page-group must have unique, relevant line numbers — no reusing or repeating sets.
                - Do not return full-page listings, only substantive testimony, objections, or statements.
                - If a page has no relevant or duplicate line numbers, omit it.
                - Do not return full-page listings — only the specific lines tied to substantive testimony, objections, or statements.
                - If incase a page has no relevant lines or same block of pagenumbers are repeating , you can omit that page from the line number dictionary.
                - Ignore if any two pages have the same array of line numbers.
            VALIDATION CHECK (you must perform before finalizing Section 1):
            1. For each page in "line_numbers", confirm every number is present on that page.
            2. Confirm no two different pages use the exact same list of line numbers.
            3. Confirm arrays are strictly increasing and 2–6 items long.
            4. If any check fails, revise the selections to comply.

                
        For every extracted chunk, include:
        - "subject": A 1-line title summarizing the main focus of these pages.
        - "content": A concise 2–3 line factual summary.
        -  "line_numbers" : A dictionary mapping page numbers to lists of line numbers that support the summary.
        Return all extracted page-group summaries in strict JSON format:
        {json_format}

        ---

        ### 2. Structured Deposition Summary
        Create a professional, litigation-ready summary organized into the following sections:  
        
        #### 1. Exhibits Table
                
        Extract all exhibits introduced or referenced in the deposition and present them in a table format:

        | Exhibit No./Name | Page Numbers | Brief Description & Relevance |
        |------------------|--------------|-------------------------------|
        | EX-1             | 12, 14, 47   | [1–2 line factual relevance]  |
        | EX-2             | 33           | [1–2 line factual relevance]  |

        Instructions:
        - Capture **every exhibit identifier** exactly as it appears (e.g., "Exhibit 12", "EX-3", "Plaintiff’s Exhibit A").  
        - Include **all page numbers** where the exhibit is either introduced, marked, or referenced in testimony.  
        - If an exhibit appears on multiple non-contiguous pages, list all page numbers separated by commas.  
        - Provide a **1–2 line neutral factual description** of the exhibit’s content or its relevance to the case.  
        - Keep it concise, litigation-ready, and fact-focused (no opinions).  
        #### 2. Legal Issue
        - Identify the primary legal issue(s) or disputes.
        - Note claims, defenses, or counterclaims.
        - Highlight if issues are contractual, statutory, regulatory, or procedural.
        - Indicate whether disputes involve interpretation of documents or factual disagreements.

        #### 3. Purpose of Deposition
        - State why this deposition was conducted.
        - Identify the strategic objective (timeline clarification, admissions, etc.).
        - Indicate type of witness (party, fact, or expert).
        - Highlight trial preparation, settlement leverage, or compliance purposes.

        #### 4. Roles
        - Name the deponent’s title and job function.
        - Explain their relevance to the case.
        - Mention if they are a decision-maker or fact witness.
        - Note other key individuals referenced.

        #### 5. Policies, Laws, or Definitions Referenced
        - List relevant policies mentioned.
        - Include applicable laws, statutes, or regulations.
        - Identify key contract clauses.
        - Note formal definitions clarified.

        #### 6. Situational Background and Key Testimony
        - Summarize critical events leading to deposition.
        - Provide chronological context.
        - Highlight crucial facts established or disputed.
        - Identify key concessions or contradictions.

        #### 7. Key Witness Statements Supporting the Case
        For each impactful or repeated statement (quoted or paraphrased in 1–2 lines), include:
        - **Speaker** — name and/or role.
        - **Situation/Context** — when and why it was said (e.g., during cross-examination, discussing an exhibit, responding to a timeline question).
        - **Impact** — concise explanation of how this strengthens the deposition’s value to the case.

       
        #### 8. Legal Recommendations
        - Suggest next litigation or discovery steps.
        - Identify additional evidence or witnesses needed.
        - Recommend motions or filings.
        - Flag risks or gaps requiring follow-up.

        ---
        
        ### General Instructions
        - Ensure the JSON summaries and the structured summary are **neutral and litigation-ready**.
        - Avoid speculation.
        - Return the final output in two sections:
        1. "Page-Group Subject Summaries (JSON)"
        2. "Structured Deposition Summary"
        """



# ============================
#  FILE UPLOAD + READ + SUMMARY
# ============================
# ============================
#  AUTO READ FILE ON UPLOAD
# ============================
if uploaded_file is not None:

    # Run only once per new upload
    if st.session_state.get("last_uploaded") != uploaded_file.name:

        st.session_state["last_uploaded"] = uploaded_file.name

        # 1️⃣ Upload file to Azure Blob
        blob_name, deposition_sas_url = upload_file_to_blob(uploaded_file)
        st.session_state["deposition_sas_url"] = deposition_sas_url
        st.session_state["blob_name"] = blob_name

        # 2️⃣ Download blob locally for text extraction
        temp_path = download_blob_to_temp(blob_name)

        # 3️⃣ Extract text from PDF or DOCX
        if temp_path.lower().endswith(".pdf"):
            text = extract_text_from_pdf(temp_path)
        else:
            text = extract_text_from_docx(temp_path)

        # Save extracted text
        st.session_state["file_text"] = text

        # 4️⃣ Show success message once
        st.success("✅ File uploaded & text extracted automatically!")

# if uploaded_file is not None:

#     # with col1:

#         # if st.button("📖 Read File"):
#         if "last_uploaded" not in st.session_state or st.session_state.last_uploaded != uploaded_file.name:
#             st.session_state.last_uploaded = uploaded_file.name

#             # Upload file ONLY here — not during upload widget
#             # blob_name = upload_file_to_blob(uploaded_file)
#             blob_name, deposition_sas_url = upload_file_to_blob(uploaded_file)
#             st.session_state["deposition_sas_url"] = deposition_sas_url
#             st.session_state["blob_name"] = blob_name

#             # Download for extraction
#             temp_path = download_blob_to_temp(blob_name)

#             if temp_path.lower().endswith(".pdf"):
#                 text = extract_text_from_pdf(temp_path)
#             else:
#                 text = extract_text_from_docx(temp_path)

#             st.session_state["file_text"] = text

#             st.success("✅ File uploaded + text extracted successfully!")


    # # --- READ FILE BUTTON ---
    # with col1:
    #     if st.button("📖 Read File"):

    #         if blob_name is None:
    #             st.error("Blob not found. Please re-upload file.")
    #             st.stop()

    #         temp_path = download_blob_to_temp(blob_name)

    #         if temp_path.lower().endswith(".pdf"):
    #             extracted = extract_text_from_pdf(temp_path)
    #         else:
    #             extracted = extract_text_from_docx(temp_path)

    #         st.session_state["file_text"] = extracted
    #         st.success("✅ File text extracted. You can now ask questions or generate summary.")

col1, col2, col3 = st.columns([1, 1, 1])

    # --- GENERATE SUMMARY BUTTON ---
with col1:
     
        if st.session_state.summary_status == "idle":

            if st.button("🧠 Generate Summary in Background"):
                blob_name = st.session_state.get("blob_name")
                if blob_name is None:
                    st.error("Please upload and read file first.")
                    st.stop()
                
                st.session_state.summary_status = "running"
                st.session_state.summary_log = []
                st.session_state.summary_error = None

                future = executor.submit(
                    background_summary,
                    blob_name,
                    api_key,
                    prompt
                )
               
                st.session_state.summary_future = future
                st.info("⚙️ Summary job started…")
                # st.rerun()

        # elif st.session_state.summary_status == "running":

        #     st.warning("⏳ Summary is being generated...")

        #     if st.session_state.summary_log:
        #         st.text("\n".join(st.session_state.summary_log[-5:]))

        #     future = st.session_state.summary_future

        #     if future and future.done():
        #         result = future.result()

        #         st.session_state.summary_log = result.get("log", [])

        #         if result.get("path"):
        #             st.session_state.summary_result = result["path"]
        #             st.session_state.summary_status = "done"
        #         else:
        #             st.session_state.summary_error = result.get("error", "Unknown error")
        #             st.session_state.summary_status = "error"

        #         st.rerun()

# with col2:
#     if st.session_state.summary_status in ("running", "done"):
#         if st.button("↻ Status", use_container_width=False):
#             if st.session_state.summary_status == "running":
#                 st.info("⏳ Summary is still in progress…")
#             else:
#                 st.rerun()
# with col2:
#     if st.session_state.summary_status == "running" or st.session_state.summary_status == "done":
#         if st.button("🔍 Check status", use_container_width=False):
#             if st.session_state.summary_status == "running":
#                 st.info("⏳ Your summary is still being inprogress..")
with col2:
    if st.session_state.summary_status in ("running", "done"):
        if st.button("🔍 Check Summary status", use_container_width=False):
            if st.session_state.summary_status == "running":
                st.session_state.show_status_msg = True
        if (
            st.session_state.summary_status == "running"
            and st.session_state.show_status_msg
        ):
            st.info("⏳ Working on your summary — explore the page in the meantime 🚀")




    # --- SHOW DOWNLOAD BUTTON ---
with col3:
        if st.session_state.summary_status == "done":
            st.success("✅ Summary generated and stored in Azure.")
            st.markdown(
                f"<a href='{st.session_state.summary_result}' target='_blank'>📄 Download Summary</a>",
                unsafe_allow_html=True
            )
            # st.markdown(f"<a href='{st.session_state.deposition_sas_url}' target='_blank'>📄 View Uploaded File</a>", unsafe_allow_html=True)


        elif st.session_state.summary_status == "error":
            st.error(f"❌ Summary failed: {st.session_state.summary_error}")

# --- Query Section ---
st.markdown("###  ⌨️Deposition Inquiry Assistant")
query_type = st.radio("Choose Input Type:", ["Dropdown", "Text Input"], horizontal=True)

st.markdown("""
    <style>
        /* 🌿 Dropdown / Multiselect Field Styling */
        div[data-baseweb="select"] {
            background-color: #e9f8ee !important;  /* light green background */
            border-radius: 10px !important;
            border: 1px solid #b6dec2 !important;
            padding: 5px !important;
            transition: all 0.2s ease-in-out;
        }

        /* Hover effect for dropdown area */
        div[data-baseweb="select"]:hover {
            background-color: #dcf5e5 !important;  /* slightly brighter green */
            border-color: #00b26d !important;
        }

        /* Selected value text */
        div[data-baseweb="select"] > div {
            color: #006b3f !important;
            font-weight: 500 !important;
        }

        /* Option list background */
        ul[role="listbox"] {
            background-color: #f5fbf6 !important;  /* dropdown open background */
            border-radius: 8px !important;
            border: 1px solid #b6dec2 !important;
        }

        /* Each dropdown item */
        li[role="option"] {
            color: #004d2c !important;
            padding: 8px 10px !important;
            border-radius: 6px !important;
        }

        /* Hover effect for each option */
        li[role="option"]:hover {
            background-color: #d8f2de !important;
            color: #003820 !important;
        }

        /* Selected tags (for multiselect) */
        div[data-baseweb="tag"] {
            background-color: #009e60 !important;
            color: white !important;
            border-radius: 12px !important;
            font-weight: 600 !important;
            padding: 4px 10px !important;
        }

        div[data-baseweb="tag"]:hover {
            background-color: #00b26d !important;
        }
    </style>
""", unsafe_allow_html=True)

if query_type == "Dropdown":
 
    depo_fields = [
    "Summarize the deposition in 5 key bullet points.",
    "List all parties, attorneys, and witnesses involved.",
    "Identify the deponent’s role and relevance to the case.",
    "What are the key issues or topics discussed in this deposition?",
    "List all exhibits referred to or marked during the deposition.",
    "Summarize the witness’s main statements related to liability.",
    "Summarize any discussions related to damages or compensation.",
    "Summarize admissions made by the deponent, if any.",
    "Identify mentions of key individuals, companies, or organizations.",
    "Summarize any clarifications or corrections made by the witness."
]

    user_input = st.multiselect("Select Deposition Fields (you can select multiple):", depo_fields)
    if len(user_input) == 0:
        st.warning("Please select at least one field.")
        st.stop()
else:
    # user_input = st.text_input("Enter your Query:")
    query =user_input= st.text_input("📝 Enter your question:")

# --- Processing Button ---
if st.button("💬 Ask AI"):
    st.session_state.pause_autorefresh = True
    text_data = st.session_state.get('file_text', '')
    
    if not text_data:
        st.warning("Please upload and read a file first.")
    else:
        user_responses = st.session_state.get('user_responses', [])
        if query_type == "Dropdown" and user_input:
            query = f"Extract the following fields: {', '.join(user_input)}"
        elif isinstance(user_input, str):
            query = user_input
        else:
            st.warning("Please enter or select a query.")
            st.stop()

        with st.spinner("Thinking... 💭"):
            response = get_chatgpt_response(query, text_data, api_key,model="gpt-4-turbo")
            user_responses.append((query, response))
            st.session_state['user_responses'] = user_responses
            st.success("✅ Response generated!")
    st.rerun()

# --- Display Responses (Chat style) ---
st.markdown("### 🧠 AI Generated Answer")
responses = st.session_state.get('user_responses', [])
if responses:
    for q, r in reversed(responses[-5:]):
        st.markdown(f"""
        <div class="response-box">
            <div class="question">🧑‍💼 <b>Question:</b> {q}</div>
            <div class="answer">🤖 <b>AI Response:</b><br>{r}</div>
        </div>
        """, unsafe_allow_html=True)
else:
    st.info("No questions yet. Upload a document and start asking!")


# --- Footer ---
st.markdown("""
<div class="footer">
    © The Wonderful Company LLC 🌳 All Rights Reserved.
</div>
""", unsafe_allow_html=True)
