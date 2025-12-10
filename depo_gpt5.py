import os
import sys
import argparse
import logging
import traceback
from docx import Document
import fitz  # PyMuPDF
import openai
# from openai import OpenAI
import requests
from docx import Document

# perplex_api_key="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" 
perplex_api_key="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"
api_key="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"


def generate_summary_with_chatgpt(text: str, prompt: str, api_key: str, model: str = "gpt-5") -> str:
    """
    Generate a summary using OpenAI ChatGPT API.
    """
    import openai
    openai.api_key = api_key
    messages = [
        {"role": "system", "content": "You are a senior legal deposition summarizer. Use the provided instructions and the deposition text to generate a clear, concise summary "},
        {"role": "user", "content": f"{prompt}\n\nContract Text:\n{text}"}
    ]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages
        # reasoning={"effort": "high"}
    )
    return response["choices"][0]["message"]["content"].strip()

# from openai import OpenAI

# def generate_summary_with_chatgpt(text: str, prompt: str, api_key: str, model: str = "gpt-4.5") -> str:
#     """
#     Generate a summary using OpenAI ChatGPT API with reasoning.
#     """
#     # client = OpenAI(api_key=api_key)
    
#     openai.api_key = api_key
#     messages = [
#         {
#             "role": "system",
#             "content": "You are a senior legal deposition summarizer. Use the provided instructions and the deposition text to generate a clear, concise summary."
#         },
#         {
#             "role": "user",
#             "content": f"{prompt}\n\nContract Text:\n{text}"
#         }
#     ]

#     # response = client.chat.completions.create(
#     #     model=model,
#     #     messages=messages,
#     #     reasoning={"effort": "high"}  # available options: "low", "medium", "high"
#     # )
#     response = openai.ChatCompletion.create(
#         model=model,
#         messages=messages,
#         temperature=0.2
#         )

#     return response.choices[0].message.content



def generate_summary_with_perplexity(text: str, prompt: str, api_key: str, model: str = "sonar-pro") -> str:
    """
    Generate a summary of a contract PDF using Perplexity's API.

    Args:
        pdf_path (str): Path to the contract PDF.
        prompt (str): Prompt to guide the summary (e.g. what fields to extract).
        api_key (str): Your Perplexity API key.
        model (str): Perplexity model to use (default: "sonar-pro").

    Returns:
        str: Summary or error message.
    """
    print(f"---------------   Generating summary with Perplexity API using model")
    try:
        # Extract text from PDF
        text = text
        if not text:
            return "(No text extracted from PDF.)"

        # Construct messages for Perplexity
        messages = [
            {
                "role": "system",
                "content": "You are a senior deposition summarizer. Use the provided instructions and the deposition text to generate a clear, concise summary .\n\n"
                    f"{prompt}\n\n deposition Text:\n{text}" ,
            },
            {
                "role": "user",
                "content": "Please generate the summary from the above deposition.",
            },
        ]

        # Send request
        response = requests.post(
            "https://api.perplexity.ai/chat/completions",
            headers={
                "Authorization": f"Bearer {perplex_api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": model,
                "messages": messages,
                "temperature": 0.2
            },
            timeout=360
        )

        if response.status_code != 200:
            logging.error(f"[Perplexity API] Error {response.status_code}: {response.text}")
            return f"(Perplexity API error {response.status_code}: {response.text})"

        data = response.json()
        if "choices" in data and data["choices"]:
            return data["choices"][0]["message"]["content"].strip()

        return "(No valid response from Perplexity API)"

    except Exception as e:
        logging.error(f"Failed to generate summary with Perplexity: {str(e)}")
        return f"Error: {str(e)}"


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logging.error(f"Failed to extract text from DOCX: {str(e)}")
        raise

# def extract_text_from_pdf(pdf_path: str) -> str:
#     """Extract text from a PDF file."""
#     try:
#         pdf_document = fitz.open(pdf_path)
#         text = "".join(page.get_text() for page in pdf_document)
#         pdf_document.close()
#         return text
#     except Exception as e:
#         logging.error(f"Failed to extract text from PDF: {str(e)}")
#         raise

def extract_text_from_pdf(pdf_path: str) -> dict:
    """
    Extract text from a PDF file, returning a dictionary with page-wise text.
    
    Args:
        pdf_path (str): Path to the PDF file.
    
    Returns:
        dict: A dictionary where keys are page numbers and values are the extracted text.
    """
    import fitz
    import logging

    page_texts = {}
    try:
        pdf_document = fitz.open(pdf_path)
        for page_number in range(len(pdf_document)):
            page = pdf_document[page_number]
            text = page.get_text("text")
            page_texts[page_number + 1] = text.strip() if text else ""
        pdf_document.close()
        return page_texts
    except Exception as e:
        logging.error(f"Failed to extract text from PDF: {str(e)}")
        raise



def extract_text_by_page_docx(docx_path: str):
    """Yield text for each 'page' in a DOCX (simulate by splitting every N paragraphs)."""
    doc = Document(docx_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Simulate pages: every 30 paragraphs = 1 page (adjust as needed)
    page_size = 30
    for i in range(0, len(paragraphs), page_size):
        yield "\n".join(paragraphs[i:i+page_size]), i // page_size + 1

def extract_text_by_page_pdf(pdf_path: str):
    """Yield text for each page in a PDF."""
    pdf_document = fitz.open(pdf_path)
    for page_num, page in enumerate(pdf_document, 1):
        yield page.get_text(), page_num
    pdf_document.close()

def generate_page_summary(page_text: str, prompt: str, api_key: str) -> str:
    """Generate a 1-2 line summary for a page using OpenAI."""
    openai.api_key = api_key
    messages = [
        {"role": "system", "content": "You are an expert legal deposition summarizer."},
        {"role": "user", "content": f"{prompt}\n\nSummarize the following page in 1-2 lines, focusing only on the most important legal or factual content:\n\n{page_text}"}
    ]
    response = openai.ChatCompletion.create(
        model="gpt-4-turbo",
        messages=messages,
        temperature=0.4
    )
    return response["choices"][0]["message"]["content"].strip()

from docx import Document

# def process_folder(input_folder: str, output_folder: str, prompt: str, api_key: str):
#     """Process all .docx and .pdf files in a folder and write summaries to output_folder."""
#     if not os.path.isdir(input_folder):
#         raise NotADirectoryError(f"Input folder not found: {input_folder}")
#     os.makedirs(output_folder, exist_ok=True)

#     files = [f for f in os.listdir(input_folder) if f.lower().endswith(('.docx', '.pdf')) and not f.startswith('~$')]
#     if not files:
#         logging.warning("No .docx or .pdf files found in the input folder.")
#         return

#     for filename in files:
#         input_path = os.path.join(input_folder, filename)
#         output_filename = os.path.splitext(filename)[0] + "_summary.docx"
#         output_path = os.path.join(output_folder, output_filename)

#         try:
#             ext = os.path.splitext(filename)[1].lower()
#             if ext == '.docx':
#                 text = extract_text_from_docx(input_path)
#             elif ext == '.pdf':
#                 text = extract_text_from_pdf(input_path)
#             else:
#                 continue

#             # Generate summary using Perplexity API
#             summary = generate_summary_with_perplexity(text, prompt, perplex_api_key)

#             # Write summary to DOCX
#             doc = Document()
#             doc.add_heading(f"Deposition Summary - {filename}", level=1)
#             doc.add_paragraph(summary)
#             doc.save(output_path)

#             logging.info(f"Summary written to {output_path}")

#         except Exception as e:
#             logging.error(f"Error processing {filename}: {str(e)}")
#             traceback.print_exc()



def process_folder(input_folder: str, output_folder: str, prompt: str, api_key: str):
    """Process all .docx and .pdf files in a folder and write summaries to output_folder."""
    if not os.path.isdir(input_folder):
        raise NotADirectoryError(f"Input folder not found: {input_folder}")
    os.makedirs(output_folder, exist_ok=True)
    # Exclude files starting with '~$'
    files = [f for f in os.listdir(input_folder) if f.lower().endswith(('.docx', '.pdf')) and not f.startswith('~$')]
    if not files:
        logging.warning("No .docx or .pdf files found in the input folder.")
        return
    for filename in files:
        input_path = os.path.join(input_folder, filename)
        output_filename = os.path.splitext(filename)[0] + ".docx"
        output_path = os.path.join(output_folder, output_filename)
        try:
            ext = os.path.splitext(filename)[1].lower()
            if ext == '.docx':
                text = extract_text_from_docx(input_path)
            elif ext == '.pdf':
                text = extract_text_from_pdf(input_path)
            else:
                continue
            
            text_py_path = output_folder + "extract_data.txt"
            print(f"Extracted text from {input_path} to {text_py_path}")
            with open(text_py_path, 'w', encoding='utf-8') as f:
                f.write(f'"""Extracted text from {input_path}"""\n\n')
                f.write(f'{text}\n')
            
            

            # summary = generate_summary_with_perplexity(text, prompt, perplex_api_key)
            summary=generate_summary_with_chatgpt(text, prompt, api_key)
            
            
            # Write summary to DOCX
            doc = Document()
            doc.add_heading(f"Deposition Summary - {filename}", level=1)
            doc.add_paragraph(summary)
            # doc.add_paragraph("Keyword Occurrences Table:\n" + keyword_table)
            doc.save(output_path)
            
            # text_py_path = output_path + "summary" + ".txt"
            # with open(text_py_path, 'w', encoding='utf-8') as f:
            #     f.write(f'"""Extracted text from {input_path}"""\n\n')
            #     f.write(f'text = """\n{text}\n"""\n')
            # # print(f"Summary for {summary} generated successfully.")
            logging.info(f"Summary written to {output_path}")
        except Exception as e:
            logging.error(f"Error processing {filename}: {str(e)}")
            traceback.print_exc()

def find_keywords_in_text_by_page(page_texts: dict, keywords: list) -> dict:
    """
    Returns a dictionary mapping each keyword to a list of page numbers where it appears.
    """
    keyword_pages = {k: [] for k in keywords}
    for page_num, text in page_texts.items():
        for k in keywords:
            if k.lower() in text.lower():
                keyword_pages[k].append(page_num)
    # Remove keywords not found
    keyword_pages = {k: v for k, v in keyword_pages.items() if v}
    return keyword_pages

def main():
    parser = argparse.ArgumentParser(description='Summarize DOCX or PDF documents or all in a folder.')
    parser.add_argument('--input', '-i', help='Input file or folder path (.docx, .pdf, or folder)', default=r"C:\Users\Teju\Desktop\ammu\955-depo\Depo_sum_sample")
    parser.add_argument('--output', '-o', help='Output summary text file or folder', default=r"C:\Users\Teju\Desktop\ammu\955-depo\Depo_sum_sample\output")
    args = parser.parse_args()
    api_key = "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"
    

    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[logging.StreamHandler(sys.stdout)]
    )
    if os.path.isdir(args.input):
        print("Files in input folder:")
        for fname in os.listdir(args.input):
            print(fname)
    
    json_format = """{
            "subject": "<short header summarizing the topic of these pages>",
            "content": "<1-2 line factual mini-summary of the testimony or events in these pages>",
            "line_numbers": {
            "<page_number>": [<only the most relevant line numbers from this page>],
            "<page_number>": [<only the most relevant line numbers from this page>]
            }
        },
        {
            "subject": "<short header summarizing the topic of these pages>",
            "content": "<1-2 line factual mini-summary of the testimony or events in these pages>",
            "line_numbers": {
            "<page_number>": [<only the most relevant line numbers from this page>],
            "<page_number>": [<only the most relevant line numbers from this page>]
            }
        }
        """

    prompt = f"""
        You are a senior legal analyst specializing in deposition analysis. Your task is to review a full deposition transcript and perform two critical functions:
        1. Page-Group Subject Summaries
        2. A structured, professional legal summary

        ---
        
        # ### 1. Page-Group Subject Summaries
            - You must review the entire deposition transcript thoroughly from start to end. 
            - Divide the transcript into sequential, non-overlapping chunks. Each chunk must:
                • Cover a continuous range of pages in order.  
                • Group together related discussions, testimony, or objections that form a coherent subject.  
                • Ensure that every page of the transcript is included in exactly one chunk (no page may be skipped or left out).  
            - For each chunk:
                • Identify a concise subject line summarizing the primary topic or testimony for that page range.  
                • Write a 2–3 line factual mini-summary of the testimony in that chunk.  
                • Keep the summary neutral, objective, and legally relevant (no opinions or speculation).  
                - Provide a **2–3 line factual mini-summary** of the content in those pages make sure all pages are included.
                - Keep it **neutral, objective, and legally relevant**.
                • Ensure that every page of the transcript is considered in sequence, but you may omit chunks if:
                    – The pages contain no substantive facts to summarize (e.g., filler, procedural headers , word glossary ).  
                    – The pages have no valid line numbers available for extraction
            - Output must cover the **entire deposition**, from the first page to the last, in properly ordered chunks.
                            
            - **Line_Numbers**:  
                - Never invent or guess line numbers. Use only numbers that truly appear on that page (available_lines or parsed from page_text).
                - Parsing: Each page is already provided as a dictionary (`1: "...", 2: "...", ...`).   — i.e., a dict mapping **page → list of 4–5 line numbers**.
                Inside each page, every `\n` corresponds to a new line number, which is marked at the start (e.g., `1`, `2`, `3`).  
                This allows you to map: **page → line numbers → text** cleanly.  
                - Select only line numbers that directly support the chunk’s subject/summary.
                - Do not include every line or filler text (e.g., "Page X", "Veritext Legal Solutions").
                - Output as a dict: "page": [line1, line2, ...].
                - Each page-group must have unique, relevant line numbers — no reusing or repeating sets.
                - Do not return full-page listings, only substantive testimony, objections, or statements.
                - If a page has no relevant or duplicate line numbers, omit it.
                - Do not return full-page listings — only the specific lines tied to substantive testimony, objections, or statements.
                - If incase a page has no relevant lines or same block of pagenumbers are repeating , you can omit that page from the line number dictionary.
                - Ignore if any two pages have the same array of line numbers.
            VALIDATION CHECK (you must perform before finalizing Section 1):
            1. For each page in "line_numbers", confirm every number is present on that page.
            2. Confirm no two different pages use the exact same list of line numbers.
            3. Confirm arrays are strictly increasing and 2–6 items long.
            4. If any check fails, revise the selections to comply.

                
        For every extracted chunk, include:
        - "subject": A 1-line title summarizing the main focus of these pages.
        - "content": A concise 2–3 line factual summary.
        -  "line_numbers" : A dictionary mapping page numbers to lists of line numbers that support the summary.
        Return all extracted page-group summaries in strict JSON format:
        {json_format}

        ---

        ### 2. Structured Deposition Summary
        Create a professional, litigation-ready summary organized into the following sections:  
        
        #### 1. Exhibits Table
                
        Extract all exhibits introduced or referenced in the deposition and present them in a table format:

        | Exhibit No./Name | Page Numbers | Brief Description & Relevance |
        |------------------|--------------|-------------------------------|
        | EX-1             | 12, 14, 47   | [1–2 line factual relevance]  |
        | EX-2             | 33           | [1–2 line factual relevance]  |

        Instructions:
        - Capture **every exhibit identifier** exactly as it appears (e.g., "Exhibit 12", "EX-3", "Plaintiff’s Exhibit A").  
        - Include **all page numbers** where the exhibit is either introduced, marked, or referenced in testimony.  
        - If an exhibit appears on multiple non-contiguous pages, list all page numbers separated by commas.  
        - Provide a **1–2 line neutral factual description** of the exhibit’s content or its relevance to the case.  
        - Keep it concise, litigation-ready, and fact-focused (no opinions).  
        #### 2. Legal Issue
        - Identify the primary legal issue(s) or disputes.
        - Note claims, defenses, or counterclaims.
        - Highlight if issues are contractual, statutory, regulatory, or procedural.
        - Indicate whether disputes involve interpretation of documents or factual disagreements.

        #### 3. Purpose of Deposition
        - State why this deposition was conducted.
        - Identify the strategic objective (timeline clarification, admissions, etc.).
        - Indicate type of witness (party, fact, or expert).
        - Highlight trial preparation, settlement leverage, or compliance purposes.

        #### 4. Roles
        - Name the deponent’s title and job function.
        - Explain their relevance to the case.
        - Mention if they are a decision-maker or fact witness.
        - Note other key individuals referenced.

        #### 5. Policies, Laws, or Definitions Referenced
        - List relevant policies mentioned.
        - Include applicable laws, statutes, or regulations.
        - Identify key contract clauses.
        - Note formal definitions clarified.

        #### 6. Situational Background and Key Testimony
        - Summarize critical events leading to deposition.
        - Provide chronological context.
        - Highlight crucial facts established or disputed.
        - Identify key concessions or contradictions.

        #### 7. Key Witness Statements Supporting the Case
        For each impactful or repeated statement (quoted or paraphrased in 1–2 lines), include:
        - **Speaker** — name and/or role.
        - **Situation/Context** — when and why it was said (e.g., during cross-examination, discussing an exhibit, responding to a timeline question).
        - **Impact** — concise explanation of how this strengthens the deposition’s value to the case.

       
        #### 8. Legal Recommendations
        - Suggest next litigation or discovery steps.
        - Identify additional evidence or witnesses needed.
        - Recommend motions or filings.
        - Flag risks or gaps requiring follow-up.

        ---
        
        ### General Instructions
        - Ensure the JSON summaries and the structured summary are **neutral and litigation-ready**.
        - Avoid speculation.
        - Return the final output in two sections:
        1. "Page-Group Subject Summaries (JSON)"
        2. "Structured Deposition Summary"
        """


    try:
        if os.path.isdir(args.input):
            process_folder(args.input, args.output, prompt, api_key)
        elif os.path.isfile(args.input):
            ext = os.path.splitext(args.input)[1].lower()
            if ext == '.docx':
                text = extract_text_from_docx(args.input)
            elif ext == '.pdf':
                text = extract_text_from_pdf(args.input)
            else:
                raise ValueError("Unsupported file type. Only .docx and .pdf are supported.")
            
            text_py_path = args.output + ".txt"
            with open(text_py_path, 'w', encoding='utf-8') as f:
                f.write(f'"""Extracted text from {args.input}"""\n\n')
                f.write(f'text = """\n{text}\n"""\n')
                    
            summary = generate_summary_with_chatgpt(text, prompt, api_key)
            # summary = summarize_document_by_page(text, prompt, api_key,args.output)
            # summary=generate_summary_with_perplexity(text, prompt, perplex_api_key)
            logging.info(f"file Summary  filee written to {args.output}")
        else:
            raise FileNotFoundError(f"Input path not found: {args.input}")
    except Exception as e:
        logging.error(f"Error: {str(e)}")
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()